"""
Document Loader Module

This module handles the loading and preprocessing of documents from various file formats:
- PDF files
- Text files
- Markdown files
- HTML files
"""

import os
import re
from typing import List, Optional
import pypdf
import markdown
from haystack import Document
from bs4 import BeautifulSoup

# 添加docx支持
try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx package not installed. DOCX file support will be disabled.")
    print("To enable DOCX support, install with: pip install python-docx")

def load_documents(directory_path: str, file_types: Optional[List[str]] = None) -> List[Document]:
    """
    Load documents from a directory, including all subdirectories recursively.
    
    Args:
        directory_path: Path to the directory containing documents
        file_types: List of file extensions to load (e.g., ['.pdf', '.txt', '.md'])
                    If None, all supported file types will be loaded
    
    Returns:
        List of Haystack Document objects
    """
    if file_types is None:
        file_types = ['.pdf', '.txt', '.md', '.html', '.htm', '.docx']
    
    documents = []
    total_files = 0
    successful_files = 0
    failed_files = 0
    encoding_success = {}
    directories_scanned = set()
    
    print(f"Scanning directory and all subdirectories: {directory_path}")
    print(f"Supported file types: {', '.join(file_types)}")
    print("This loader supports Chinese documents with multiple encodings (UTF-8, GB18030, GBK, GB2312, BIG5)")
    
    # Count total files first and gather directory information
    for root, _, files in os.walk(directory_path):
        rel_path = os.path.relpath(root, directory_path)
        if rel_path != '.':
            directories_scanned.add(rel_path)
        
        for file in files:
            if any(file.lower().endswith(ft) for ft in file_types):
                total_files += 1
    
    if total_files == 0:
        print(f"No supported files found in {directory_path} or its subdirectories")
        return []
    
    subdirs_str = f" in {len(directories_scanned)} subdirectories" if directories_scanned else ""
    print(f"Found {total_files} supported files to process{subdirs_str}")
    
    # 显示扫描到的子目录
    if directories_scanned:
        print("Subdirectories that will be processed:")
        for subdir in sorted(directories_scanned):
            print(f"  - {subdir}")
    
    # 按照目录层级加载文档
    for root, _, files in os.walk(directory_path):
        # 显示当前处理的目录
        rel_path = os.path.relpath(root, directory_path)
        dir_display = f"{rel_path}/" if rel_path != '.' else "(root)"
        
        # 检查当前目录是否有支持的文件
        has_supported_files = any(any(file.lower().endswith(ft) for ft in file_types) for file in files)
        if has_supported_files:
            print(f"\nProcessing directory: {dir_display}")
        
        for file in files:
            if any(file.lower().endswith(ft) for ft in file_types):
                file_path = os.path.join(root, file)
                try:
                    # 从文件名中提取标题（去除文件扩展名）
                    filename_without_ext = os.path.splitext(file)[0]
                    
                    if file.lower().endswith('.pdf'):
                        docs = load_pdf(file_path, filename_without_ext)
                    elif file.lower().endswith('.txt'):
                        docs = load_text(file_path, filename_without_ext)
                    elif file.lower().endswith('.md'):
                        docs = load_markdown(file_path, filename_without_ext)
                    elif file.lower().endswith(('.html', '.htm')):
                        docs = load_html(file_path, filename_without_ext)
                    elif file.lower().endswith('.docx'):
                        docs = load_docx(file_path, filename_without_ext)
                    else:
                        continue
                    
                    if docs:
                        successful_files += 1
                        documents.extend(docs)
                        # Record which encoding succeeded (only for text-based files)
                        if hasattr(docs[0], 'meta') and 'encoding' in docs[0].meta:
                            encoding = docs[0].meta['encoding']
                            encoding_success[encoding] = encoding_success.get(encoding, 0) + 1
                        print(f"  Loaded {len(docs)} segments from {file}")
                    else:
                        failed_files += 1
                        print(f"  No content extracted from {file}")
                except Exception as e:
                    failed_files += 1
                    print(f"  Error loading {file}: {e}")
    
    print(f"\nDocument loading summary:")
    print(f"Total files processed: {total_files}")
    print(f"Successfully loaded files: {successful_files}")
    print(f"Failed files: {failed_files}")
    print(f"Total document segments created: {len(documents)}")
    
    if encoding_success:
        print("\nEncoding statistics for text files:")
        for encoding, count in encoding_success.items():
            print(f"  - {encoding}: {count} files")
    
    return documents

def load_pdf(file_path: str, title: Optional[str] = None) -> List[Document]:
    """
    Load and preprocess PDF documents.
    
    Args:
        file_path: Path to the PDF file
        title: Document title (typically from filename)
    
    Returns:
        List of Document objects
    """
    documents = []
    
    try:
        pdf_reader = pypdf.PdfReader(file_path)
        for i, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():  # Skip empty pages
                    # Sometimes PDF text extraction may introduce encoding issues
                    # Try to clean the text
                    try:
                        # Handle potential encoding issues by decoding and re-encoding
                        text = text.encode('utf-8', errors='ignore').decode('utf-8')
                    except:
                        pass  # Use original text if re-encoding fails
                    
                    doc = Document(
                        content=text,
                        meta={
                            "source": file_path,
                            "page": i + 1,
                            "file_type": "pdf",
                            "title": title
                        }
                    )
                    documents.append(doc)
            except Exception as page_error:
                print(f"Error extracting text from page {i+1} in PDF {file_path}: {page_error}")
                # Continue to the next page even if one fails
    except Exception as e:
        print(f"Error processing PDF {file_path}: {e}")
    
    return documents

def load_text(file_path: str, title: Optional[str] = None) -> List[Document]:
    """
    Load and preprocess text documents.
    
    Args:
        file_path: Path to the text file
        title: Document title (typically from filename)
    
    Returns:
        List of Document objects
    """
    # Try different encodings for Chinese text files
    encodings = ['utf-8', 'gb18030', 'gbk', 'gb2312', 'big5']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                text = file.read()
            
            doc = Document(
                content=text,
                meta={
                    "source": file_path,
                    "file_type": "txt",
                    "encoding": encoding,
                    "title": title
                }
            )
            return [doc]
        except UnicodeDecodeError:
            # Try next encoding
            continue
        except Exception as e:
            print(f"Error processing text file {file_path}: {e}")
            return []
    
    print(f"Error processing text file {file_path}: Unable to decode with any supported encoding")
    return []

def load_markdown(file_path: str, title: Optional[str] = None) -> List[Document]:
    """
    Load and preprocess Markdown documents.
    
    Args:
        file_path: Path to the Markdown file
        title: Document title (typically from filename)
    
    Returns:
        List of Document objects
    """
    # Try different encodings for Chinese text files
    encodings = ['utf-8', 'gb18030', 'gbk', 'gb2312', 'big5']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                md_text = file.read()
            
            # Convert Markdown to HTML then strip HTML tags to get plain text
            html = markdown.markdown(md_text)
            text = re.sub('<[^<]+?>', '', html)
            
            doc = Document(
                content=text,
                meta={
                    "source": file_path,
                    "file_type": "md",
                    "encoding": encoding,
                    "title": title
                }
            )
            return [doc]
        except UnicodeDecodeError:
            # Try next encoding
            continue
        except Exception as e:
            print(f"Error processing Markdown file {file_path}: {e}")
            return []
    
    print(f"Error processing Markdown file {file_path}: Unable to decode with any supported encoding")
    return []

def load_html(file_path: str, title: Optional[str] = None) -> List[Document]:
    """
    Load and preprocess HTML documents.
    
    Args:
        file_path: Path to the HTML file
        title: Document title (typically from filename)
    
    Returns:
        List of Document objects
    """
    # Try different encodings for Chinese HTML files
    encodings = ['utf-8', 'gb18030', 'gbk', 'gb2312', 'big5']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                html_content = file.read()
            
            # Parse HTML and extract text
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Get text
            text = soup.get_text(separator=' ', strip=True)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            if text:
                # 首选<title>标签中的标题，其次是传入的文件名标题
                html_title = soup.title.string if soup.title else None
                doc_title = html_title or title
                
                doc = Document(
                    content=text,
                    meta={
                        "source": file_path,
                        "file_type": "html",
                        "encoding": encoding,
                        "title": doc_title
                    }
                )
                return [doc]
        except UnicodeDecodeError:
            # Try next encoding
            continue
        except Exception as e:
            print(f"Error processing HTML file {file_path}: {e}")
            return []
    
    print(f"Error processing HTML file {file_path}: Unable to decode with any supported encoding")
    return []

def load_docx(file_path: str, title: Optional[str] = None) -> List[Document]:
    """
    Load and preprocess DOCX documents.
    
    Args:
        file_path: Path to the DOCX file
        title: Document title (typically from filename)
    
    Returns:
        List of Document objects
    """
    if not DOCX_SUPPORT:
        print(f"DOCX support is disabled. Cannot process {file_path}")
        return []
    
    documents = []
    
    try:
        doc = docx.Document(file_path)
        
        # 处理普通文档内容
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # 处理表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        text = "\n".join(full_text)
        
        if text.strip():
            doc = Document(
                content=text,
                meta={
                    "source": file_path,
                    "file_type": "docx",
                    "title": title
                }
            )
            documents.append(doc)
        
    except Exception as e:
        print(f"Error processing DOCX {file_path}: {e}")
    
    return documents

def chunk_documents(documents: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """
    Split documents into smaller chunks for better processing.
    
    Args:
        documents: List of Document objects
        chunk_size: Maximum size of each chunk
        chunk_overlap: Overlap between chunks
    
    Returns:
        List of chunked Document objects
    """
    chunked_documents = []
    
    for doc in documents:
        text = doc.content
        
        # If the document is already small enough, keep it as is
        if len(text) <= chunk_size:
            chunked_documents.append(doc)
            continue
        
        # Split into chunks
        start = 0
        chunk_id = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to find a good break point (newline or period)
            if end < len(text):
                # Look for newline or period around the chunk_size boundary
                possible_end = text.rfind('\n', start, end)
                if possible_end == -1 or (end - possible_end) > chunk_size / 4:
                    possible_end = text.rfind('. ', start, end)
                
                if possible_end != -1 and (end - possible_end) < chunk_size / 4:
                    end = possible_end + 2  # Include the period and space
            
            chunk_text = text[start:end]
            
            if chunk_text.strip():  # Skip empty chunks
                chunk_meta = doc.meta.copy() if doc.meta else {}
                chunk_meta["chunk_id"] = chunk_id
                
                chunked_doc = Document(
                    content=chunk_text,
                    meta=chunk_meta
                )
                chunked_documents.append(chunked_doc)
                
                chunk_id += 1
            
            start = end - chunk_overlap
    
    return chunked_documents

def is_duplicate_document(document: Document, document_store, threshold: float = 0.95) -> bool:
    """
    Check if a document already exists in the document store by comparing source metadata.
    
    Args:
        document: Document to check
        document_store: Haystack document store
        threshold: Similarity threshold for considering a document as duplicate
    
    Returns:
        True if the document appears to be a duplicate, False otherwise
    """
    # First check: If the document has 'source' in meta, look for exact source match
    if document.meta and "source" in document.meta:
        source = document.meta["source"]
        
        print(f"Checking if document with source '{source}' already exists...")
        
        try:
            # 直接在文档的元数据中查找匹配的 source
            # 注意：不同的文档存储实现查询方式可能不同
            # 这里先尝试特定的查询方式
            if hasattr(document_store, "_collection"):
                # ChromaDocumentStore特定的查询方式
                print("使用ChromaDocumentStore特定的查询方式...")
                try:
                    # 获取文档内容，包括元数据
                    result = document_store._collection.get(include=["metadatas"])
                    
                    if result and "metadatas" in result and result["metadatas"]:
                        metadatas = result["metadatas"]
                        print(f"Found {len(metadatas)} documents with metadata")
                        
                        for meta in metadatas:
                            if meta and "source" in meta and meta["source"] == source:
                                print(f"DUPLICATE DETECTED: Document with source '{source}' already exists")
                                return True
                        
                        # 打印一些元数据样本以帮助调试
                        if metadatas and len(metadatas) > 0:
                            print(f"Sample metadata: {metadatas[0]}")
                    else:
                        print("No documents with metadata found in the collection")
                except Exception as e:
                    print(f"Error accessing ChromaDB collection directly: {e}")
            
            # 标准方式：使用过滤器
            all_docs = document_store.filter_documents({})
            print(f"Standard query found {len(all_docs)} documents")
            
            for existing_doc in all_docs:
                if existing_doc.meta and "source" in existing_doc.meta:
                    existing_source = existing_doc.meta["source"]
                    print(f"Comparing with existing source: '{existing_source}'")
                    if existing_source == source:
                        print(f"DUPLICATE DETECTED: Document with source '{source}' already exists in the document store")
                        return True
        except Exception as e:
            print(f"Error checking for duplicates: {e}")
    
    print("No duplicate detected based on source metadata")
    return False 